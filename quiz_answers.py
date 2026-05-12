from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── Style setup ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ═══════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════
title = doc.add_heading('Compiler Theory Quiz — Answer Key', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ═══════════════════════════════════════════
# Q1
# ═══════════════════════════════════════════
doc.add_heading('Question 1 — Compiler Front End vs. Middle/Rear End  (10 pts)', level=2)

p = doc.add_paragraph()
p.add_run('Question: ').bold = True
p.add_run('Which compiler phase is not "front end" phase?')

doc.add_paragraph('Options:', style='List Bullet')
opts = [
    'Lexical analysis',
    'Parsing',
    'Semantic analysis',
    'Optimization'
]
for o in opts:
    doc.add_paragraph(o, style='List Bullet 2')

ans = doc.add_paragraph()
ans.add_run('Correct Answer: ').bold = True
ans.add_run('d) Optimization')
ans.runs[-1].bold = True
ans.runs[-1].font.color.rgb = RGBColor(0, 100, 0)

explanation = doc.add_paragraph()
explanation.add_run('Explanation: ').bold = True
explanation.add_run(
    'The front end of a compiler consists of lexical analysis, parsing, and semantic analysis. '
    'Optimization belongs to the middle end (or sometimes rear end) of the compiler pipeline, '
    'after the front end has produced an intermediate representation (IR). Code generation '
    'is the final stage.'
)

# ═══════════════════════════════════════════
# Q2
# ═══════════════════════════════════════════
doc.add_heading('Question 2 — Type Environments and Method Signatures  (10 pts)', level=2)

p = doc.add_paragraph()
p.add_run('Question: ').bold = True
p.add_run('Which type environment should we use when we want to know the method signatures?')

doc.add_paragraph('Options:', style='List Bullet')
for o in ['O', 'M', 'C']:
    doc.add_paragraph(o, style='List Bullet 2')

ans = doc.add_paragraph()
ans.add_run('Correct Answer: ').bold = True
ans.add_run('c) C')
ans.runs[-1].bold = True
ans.runs[-1].font.color.rgb = RGBColor(0, 100, 0)

explanation = doc.add_paragraph()
explanation.add_run('Explanation: ').bold = True
explanation.add_run(
    'In type-system environments for object-oriented languages:\n'
    '  •  C (Class Environment) — stores class declarations, their fields, and method signatures.\n'
    '  •  M (Method/Module Environment) — stores method bindings / bodies.\n'
    '  •  O (Object Environment) — stores object instances / field values.\n\n'
    'Method signatures (the interface — return type and parameter types) are resolved through '
    'the class environment C, since classes define which methods exist and what their signatures are. '
    'When type-checking a method call like obj.method(args), the compiler looks up obj\'s class in C '
    'to verify the method exists and that args match the signature.'
)

# ═══════════════════════════════════════════
# Q3
# ═══════════════════════════════════════════
doc.add_heading('Question 3 — Free vs. Bound Variables  (10 pts)', level=2)

p = doc.add_paragraph()
p.add_run('Question: ').bold = True
p.add_run('Which variable is not free variable?')

doc.add_paragraph('Options:', style='List Bullet')
opt3 = [
    'x',
    'x + y',
    '"In Let y: int ← 4 in x+y: The occurrence of y"',
    '"In Let y: int ← 4 in x+y: The occurrence of x"'
]
for o in opt3:
    doc.add_paragraph(o, style='List Bullet 2')

ans = doc.add_paragraph()
ans.add_run('Correct Answer: ').bold = True
ans.add_run('c) "In Let y: int ← 4 in x+y: The occurrence of y"')
ans.runs[-1].bold = True
ans.runs[-1].font.color.rgb = RGBColor(0, 100, 0)

explanation = doc.add_paragraph()
explanation.add_run('Explanation: ').bold = True
explanation.add_run(
    'A free variable is one whose value is not bound within the given scope.\n\n'
    'In the expression "Let y: int ← 4 in x + y":\n'
    '  •  y is bound by the let binding (y ← 4), so y is NOT free at its occurrence.\n'
    '  •  x is not bound anywhere in the expression, so x IS free.\n\n'
    'Therefore, the occurrence of y is not a free variable (it is a bound variable occurrence). '
    'This makes option (c) the correct answer — y is bound, not free.\n\n'
    'For reference:\n'
    '  •  Option (a) "x" alone — x is free (no binding).\n'
    '  •  Option (b) "x + y" — x is free, y is bound; the set of free vars is {x}.\n'
    '  •  Option (d) "The occurrence of x" — x is free.'
)

# ═══════════════════════════════════════════
# Q4
# ═══════════════════════════════════════════
doc.add_heading('Question 4 — Recognizing Viable Prefixes  (10 pts)', level=2)

p = doc.add_paragraph()
p.add_run('Question: ').bold = True
p.add_run(
    'As shown in "Recognizing Viable Prefixes," given the grammar and the incomplete NFA:'
)

# Grammar block
doc.add_heading('Grammar:', level=3)
grammar = [
    "S' → E",
    "E → T + E | T",
    "T → int * T | int | (E)"
]
for g in grammar:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('NFA (Recognizing Viable Prefixes):', level=3)

# Add a table for the NFA states
table = doc.add_table(rows=1, cols=3)
table.style = 'Light List Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'State'
hdr_cells[1].text = 'Item'
hdr_cells[2].text = 'Transitions On'

nfa_data = [
    ('q₀', "S' → ·E", 'E → q₁'),
    ('q₁', "S' → E·", 'accepting state'),
    ('q₂', "E → ·T", 'T → q₃  |  + → q₄'),
    ('q₃', "T → ·int * T", 'int → q₅  |  ( → q₆'),
    ('q₄', "E → T + ·E", 'E → q₇'),
    ('q₅', "T → int · * T", '* → q₈  |  (end of T)'),
    ('q₆', "T → ( ·E )", 'E → q₉'),
]

for state, item, trans in nfa_data:
    row_cells = table.add_row().cells
    row_cells[0].text = state
    row_cells[1].text = item
    row_cells[2].text = trans

ans = doc.add_paragraph()
ans.add_run('Correct Answer: ').bold = True
ans.add_run(
    'The viable prefix property: A viable prefix is a prefix of a right-sentential form '
    'that does not continue past the right end of the rightmost handle of that sentential form.\n\n'
    'In LR(1) parsing, viable prefixes are recognized by a DFA constructed from an NFA. '
    'The NFA states correspond to LR(0) items (grammar productions with a "dot" marking '
    'how much has been parsed). Key closure and goto operations build the state machine.'
)

explanation = doc.add_paragraph()
explanation.add_run('Key Concepts:').bold = True

concepts = [
    'CLOSURE: If item A → α·Bβ is in a state, add B → ·γ for every production B → γ.',
    'GOTO: Transition on symbol X moves the dot past X: A → α·Xβ becomes A → αX·β.',
    'The initial state contains S\' → ·E and all items reachable via closure on E.',
    'From the start state, reading E leads to S\' → E· (accepting state — input is valid).',
    'Reading T from the closure leads to E → T· (shift-reduce: reduce by E → T).',
    'Reading int leads to T → int· (shift-reduce: reduce by T → int).',
    'Reading "(" leads to T → (·E), requiring recursive parsing of E inside parens.'
]
for c in concepts:
    doc.add_paragraph(c, style='List Bullet')

explanation2 = doc.add_paragraph()
explanation2.add_run('Note: ').bold = True
explanation2.add_run(
    'The exact answer depends on what the question is asking (which item is missing, '
    'which transition is incorrect, etc.). The NFA above captures the standard viable-prefix '
    'recognizer for this grammar. In LR parsing, every state represents a set of items '
    'corresponding to a viable prefix, and the parser uses goto transitions to move between states.'
)

# ═══════════════════════════════════════════
# SUMMARY TABLE
# ═══════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Answer Summary', level=1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light List Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Question'
hdr_cells[1].text = 'Answer'
hdr_cells[2].text = 'Key Concept'

summary = [
    ('Q1 — Front End Phase', 'd) Optimization', 'Compiler pipeline stages'),
    ('Q2 — Method Signatures', 'c) C', 'Type environment hierarchy'),
    ('Q3 — Not a Free Variable', 'c) The occurrence of y', 'Free vs. bound variable binding'),
    ('Q4 — Viable Prefixes', 'NFA-based LR item states', 'Viable prefix recognition & closure/goto'),
]
for q, a, k in summary:
    row_cells = table.add_row().cells
    row_cells[0].text = q
    row_cells[1].text = a
    row_cells[2].text = k

# ── Save ──
out_path = '/workspace/compiler/quiz_answers.docx'
doc.save(out_path)
print(f'Document saved to {out_path}')
